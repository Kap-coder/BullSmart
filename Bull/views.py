
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_GET
from django.urls import reverse
from Bull.models import BulletinTemplate, SchoolYear, Term, Sequence, Classroom, Student, Subject, Grade, Bulletin, ClassSubject, Teacher
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse, HttpResponseForbidden, FileResponse, Http404, HttpResponseRedirect
from Bull.templatetags import bulletin_tags
from django.db import models
from django.shortcuts import get_object_or_404
from django.contrib import messages
from django.db.models import Count, Q
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth import get_user_model
from django.core.paginator import Paginator
from django import forms
from django.conf import settings
from Bull.forms import StudentForm, ImportStudentsForm, ExportStudentsForm, TeacherForm, BulletinTemplateForm
import openpyxl
import io
import os
import openpyxl
from zipfile import BadZipFile

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas




#--------------------------------------------------
# authentification et BASE
#---------------------------------------------------




# Modèles (imports complémentaires utilisés dans les fonctions)
from Bull.models import StudentSubject, User

# Fonctions utilitaires
def is_admin_or_secretary(user):
    return user.is_authenticated and (user.role == 'admin' or user.role == 'secretary')

def user_can_manage_cs(user, cs):
    # Admins and secretaries can always manage
    if not user.is_authenticated:
        return False
    if getattr(user, 'role', None) in ('admin', 'secretary'):
        return True
    # Teachers: allow if they are assigned to the ClassSubject or head_teacher of the class
    try:
        teacher = getattr(user, 'teacher', None)
        if teacher is None:
            return False
        if cs.teacher and cs.teacher.id == teacher.id:
            return True
        if cs.classroom and cs.classroom.head_teacher and cs.classroom.head_teacher.id == teacher.id:
            return True
    except Exception:
        return False
    return False

# Pages publiques
def home(request):
    return render(request, 'Bull/home.html')

def login_view(request):
    if request.user.is_authenticated:
        return redirect('profile')
    error = False
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('profile')
        else:
            error = True
    return render(request, 'Bull/login.html', {'form': {}, 'error': error})

def logout_view(request):
    logout(request)
    return redirect('login')


#-------------------------------------------------
# fin AUTH et BASE
#-------------------------------------------------

# --------------------------------------------------
#    GESTIONS DES CLASSE
# ---------------------------------------------------


@login_required
@user_passes_test(is_admin_or_secretary)
def classes_view(request):
    classes = Classroom.objects.select_related('head_teacher').all()
    return render(request, 'Bull/classes.html', {'classes': classes})

@login_required
@user_passes_test(is_admin_or_secretary)
def class_add_view(request):
    from .models import Teacher
    teachers = Teacher.objects.select_related('user').all()
    if request.method == 'POST':
        name = request.POST.get('name')
        level = request.POST.get('level')
        series = request.POST.get('series')
        option = request.POST.get('option')
        head_teacher_id = request.POST.get('head_teacher')
        if name and level and option:
            if Classroom.objects.filter(name=name, level=level, option=option).exists():
                messages.error(request, "Cette classe existe déjà.")
            else:
                head_teacher = Teacher.objects.filter(id=head_teacher_id).first() if head_teacher_id else None
                Classroom.objects.create(
                    name=name,
                    level=level,
                    series=series,
                    option=option,
                    head_teacher=head_teacher
                )
                messages.success(request, "Classe ajoutée avec succès.")
                return redirect('classes')
        else:
            messages.error(request, "Veuillez remplir tous les champs obligatoires.")
    return render(request, 'Bull/class_add.html', {'teachers': teachers})

@login_required
@user_passes_test(is_admin_or_secretary)
def class_detail_view(request, class_id):
    classroom = get_object_or_404(Classroom, id=class_id)
    from Bull.models import Student
    students = Student.objects.filter(classroom=classroom)
    total = students.count()
    filles = students.filter(gender='F').count()
    garcons = students.filter(gender='M').count()
    redoublants = students.filter(repeater=True).count()
    stats = {
        'total': total,
        'filles': filles,
        'garcons': garcons,
        'redoublants': redoublants,
    }
    return render(request, 'Bull/class_detail.html', {'classroom': classroom, 'stats': stats})

class ClassroomForm(forms.ModelForm):
    class Meta:
        model = Classroom
        fields = ['name', 'level', 'series', 'head_teacher']

@login_required
@user_passes_test(is_admin_or_secretary)
def class_edit_view(request, class_id):
    classroom = Classroom.objects.get(id=class_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST, instance=classroom)
        if form.is_valid():
            form.save()
            return redirect('classes')
    else:
        form = ClassroomForm(instance=classroom)
    return render(request, 'Bull/class_edit.html', {'form': form, 'classroom': classroom})



# ---------------------------------------------
#     GESTIONS DE MATIERES
# ---------------------------------------------


@login_required
@user_passes_test(is_admin_or_secretary)
def subjects_view(request):
    from Bull.models import Subject
    subjects = Subject.objects.all()
    return render(request, 'Bull/subjects.html', {'subjects': subjects})

@login_required
@user_passes_test(is_admin_or_secretary)
def subject_add_view(request):
    from Bull.models import Subject
    if request.method == 'POST':
        code = request.POST.get('code')
        name = request.POST.get('name')
        category = request.POST.get('category')
        if code and name and category:
            if Subject.objects.filter(code=code).exists():
                messages.error(request, "Ce code existe déjà.")
            else:
                Subject.objects.create(code=code, name=name, category=category)
                messages.success(request, "Matière ajoutée avec succès.")
                return redirect('subjects')
        else:
            messages.error(request, "Veuillez remplir tous les champs.")
    return render(request, 'Bull/subject_add.html')

@login_required
@user_passes_test(is_admin_or_secretary)
def subject_edit_view(request, subject_id):
    from Bull.models import Subject
    subject = Subject.objects.get(id=subject_id)
    if request.method == 'POST':
        code = request.POST.get('code')
        name = request.POST.get('name')
        category = request.POST.get('category')
        if code and name and category:
            if Subject.objects.filter(code=code).exclude(id=subject_id).exists():
                messages.error(request, "Ce code existe déjà.")
            else:
                subject.code = code
                subject.name = name
                subject.category = category
                subject.save()
                messages.success(request, "Matière modifiée avec succès.")
                return redirect('subjects')
        else:
            messages.error(request, "Veuillez remplir tous les champs.")
    return render(request, 'Bull/subject_edit.html', {'subject': subject})

@login_required
@user_passes_test(is_admin_or_secretary)
def subject_detail_view(request, subject_id):
    from Bull.models import Subject, ClassSubject, Classroom, Teacher
    subject = Subject.objects.get(id=subject_id)
    classrooms = Classroom.objects.all()
    teachers = Teacher.objects.select_related('user').all()

    # Filtrage et recherche
    filter_class = request.GET.get('class')
    search = request.GET.get('search', '').strip()
    class_subjects = ClassSubject.objects.filter(subject=subject).select_related('classroom', 'teacher')
    if filter_class:
        class_subjects = class_subjects.filter(classroom__id=filter_class)
    if search:
        class_subjects = class_subjects.filter(
            models.Q(classroom__name__icontains=search) |
            models.Q(teacher__user__first_name__icontains=search) |
            models.Q(teacher__user__last_name__icontains=search)
        )

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(class_subjects, 10)  # 10 par page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'Bull/subject_detail.html', {
        'subject': subject,
        'class_subjects': page_obj,
        'classrooms': classrooms,
        'teachers': teachers,
        'filter_class': filter_class,
        'search': search,
        'page_obj': page_obj
    })

@login_required
@user_passes_test(is_admin_or_secretary)
def subject_card_list_view(request):
    subjects = Subject.objects.all()
    return render(request, 'Bull/subject_card_list.html', {'subjects': subjects})

@login_required
@user_passes_test(is_admin_or_secretary)
def subject_class_cards_view(request, subject_id):
    subject = Subject.objects.get(id=subject_id)
    classsubjects = ClassSubject.objects.filter(subject=subject).select_related('classroom', 'teacher')
    # debug flag via ?debug=1
    debug = request.GET.get('debug') == '1'
    # Calculer les stats pour chaque classe associée
    class_cards = []
    from Bull.models import Student
    for cs in classsubjects:
        # queryset par filtre explicit
        students_qs = Student.objects.filter(classroom=cs.classroom)
        total = students_qs.count()
        filles = students_qs.filter(gender='F').count()
        garcons = students_qs.filter(gender='M').count()
        redoublants = students_qs.filter(repeater=True).count()
        # reverse relation (via related_name) — peut échouer si relation non configurée
        try:
            reverse_count = cs.classroom.students.count()
        except Exception:
            reverse_count = None
        card = {
            'cs': cs,
            'students': students_qs,
            'stats': {
                'total': total,
                'filles': filles,
                'garcons': garcons,
                'redoublants': redoublants,
            }
        }
        if debug:
            card['debug'] = {
                'classroom_id': getattr(cs.classroom, 'id', None),
                'count_via_filter': total,
                'count_via_reverse': reverse_count,
            }
        class_cards.append(card)
    return render(request, 'Bull/subject_class_cards.html', {
        'subject': subject,
        'class_cards': class_cards,
        'debug': debug,
    })



# ----------------------------------------------------
#    gestion des associations classe-matière
# ----------------------------------------------------


@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_add_view(request, subject_id):
    subject = Subject.objects.get(id=subject_id)
    # Filtrer pour n'afficher que les classes non associées à ce sujet
    associated_classroom_ids = ClassSubject.objects.filter(subject=subject).values_list('classroom_id', flat=True)
    classrooms = Classroom.objects.exclude(id__in=associated_classroom_ids)
    teachers = Teacher.objects.select_related('user').all()
    if request.method == 'POST':
        classroom_id = request.POST.get('classroom')
        teacher_id = request.POST.get('teacher')
        coefficient = request.POST.get('coefficient')
        if classroom_id and coefficient:
            classroom = Classroom.objects.get(id=classroom_id)
            teacher = Teacher.objects.get(id=teacher_id) if teacher_id else None
            if ClassSubject.objects.filter(classroom=classroom, subject=subject).exists():
                return render(request, 'Bull/classsubject_add.html', {
                    'subject': subject,
                    'classrooms': classrooms,
                    'teachers': teachers,
                    'error': "Cette association existe déjà."
                })
            ClassSubject.objects.create(classroom=classroom, subject=subject, teacher=teacher, coefficient=coefficient)
            return redirect('subject_detail', subject_id=subject_id)
        else:
            return render(request, 'Bull/classsubject_add.html', {
                'subject': subject,
                'classrooms': classrooms,
                'teachers': teachers,
                'error': "Veuillez remplir tous les champs."
            })
    return render(request, 'Bull/classsubject_add.html', {
        'subject': subject,
        'classrooms': classrooms,
        'teachers': teachers
    })
    return redirect('subject_detail', subject_id=subject_id)

@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_edit_view(request, cs_id):
    cs = ClassSubject.objects.get(id=cs_id)
    teachers = Teacher.objects.select_related('user').all()
    if request.method == 'POST':
        coefficient = request.POST.get('coefficient')
        teacher_id = request.POST.get('teacher')
        if coefficient:
            cs.coefficient = coefficient
            cs.teacher = Teacher.objects.get(id=teacher_id) if teacher_id else None
            cs.save()
            messages.success(request, "Association modifiée.")
        else:
            messages.error(request, "Veuillez remplir tous les champs.")
        return redirect('subject_detail', subject_id=cs.subject.id)
    return redirect('subject_detail', subject_id=cs.subject.id)

@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_delete_view(request, cs_id):
    cs = ClassSubject.objects.get(id=cs_id)
    subject_id = cs.subject.id
    if request.method == 'POST':
        cs.delete()
        messages.success(request, "Association supprimée.")
        return redirect('subject_detail', subject_id=subject_id)
    return redirect('subject_detail', subject_id=subject_id)

@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_detail_page(request, cs_id):
    from Bull.models import ClassSubject
    classsubject = ClassSubject.objects.select_related('subject', 'classroom', 'teacher').get(id=cs_id)
    return render(request, 'Bull/classsubject_detail.html', {'classsubject': classsubject})

@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_edit_page(request, cs_id):
    from Bull.models import ClassSubject, Teacher
    classsubject = ClassSubject.objects.select_related('teacher').get(id=cs_id)
    teachers = Teacher.objects.select_related('user').all()
    if request.method == 'POST':
        coefficient = request.POST.get('coefficient')
        teacher_id = request.POST.get('teacher')
        if coefficient:
            classsubject.coefficient = coefficient
            classsubject.teacher = Teacher.objects.get(id=teacher_id) if teacher_id else None
            classsubject.save()
            return redirect('classsubject_detail', cs_id=classsubject.id)
    return render(request, 'Bull/classsubject_edit.html', {'classsubject': classsubject, 'teachers': teachers})

@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_delete_page(request, cs_id):
    from Bull.models import ClassSubject
    classsubject = ClassSubject.objects.get(id=cs_id)
    if request.method == 'POST':
        subject_id = classsubject.subject.id
        classsubject.delete()
        return redirect('subject_detail', subject_id=subject_id)
    return render(request, 'Bull/classsubject_delete.html', {'classsubject': classsubject})




# ------------------------------------------
# gestion des notes
# -----------------------------------------


@login_required
@user_passes_test(is_admin_or_secretary)
def classsubject_students_view(request, cs_id):
    cs = ClassSubject.objects.select_related('classroom', 'subject').get(id=cs_id)
    students = cs.classroom.students.all().order_by('last_name', 'first_name')
    # from Bull.models import Grade, Sequence, Term, SchoolYear
    schoolyears = SchoolYear.objects.all()
    terms = Term.objects.all()
    sequences = Sequence.objects.all()
    # Récupère les messages transmis par la redirection
    message = request.GET.get('success')
    error_message = request.GET.get('error')
    selected_schoolyear_id = request.GET.get('schoolyear')
    selected_term_id = request.GET.get('term')
    selected_sequence_id = request.GET.get('sequence')

    # If not selected, use active sequence/term/year
    if not (selected_schoolyear_id and selected_term_id and selected_sequence_id):
        active_seq = Sequence.objects.filter(active=True).select_related('term__school_year').first()
        if active_seq:
            selected_sequence_id = str(active_seq.id)
            selected_term_id = str(active_seq.term.id)
            selected_schoolyear_id = str(active_seq.term.school_year.id)

    # Enforce only one active sequence per term/year
    if selected_sequence_id:
        active_seqs = Sequence.objects.filter(term_id=selected_term_id, active=True)
        if active_seqs.count() > 1:
            # Keep only the first as active
            for seq in active_seqs[1:]:
                seq.active = False
                seq.save()
    grades = None
    selected_schoolyear_obj = selected_term_obj = selected_sequence_obj = None
    # Only show a list when a school year is selected (and term + sequence)
    if selected_schoolyear_id:
        try:
            selected_schoolyear_obj = SchoolYear.objects.get(id=selected_schoolyear_id)
        except SchoolYear.DoesNotExist:
            selected_schoolyear_obj = None
        if selected_term_id and selected_sequence_id:
            try:
                selected_term_obj = Term.objects.get(id=selected_term_id)
            except Term.DoesNotExist:
                selected_term_obj = None
            try:
                selected_sequence_obj = Sequence.objects.get(id=selected_sequence_id)
            except Sequence.DoesNotExist:
                selected_sequence_obj = None

            # Ensure StudentSubject and Grade records exist for every student in the class for this sequence/term
            from Bull.models import StudentSubject
            created_ss = 0
            created_grades = 0
            for student in students:
                try:
                    ss, ss_new = StudentSubject.objects.get_or_create(
                        student=student,
                        subject=cs.subject,
                        defaults={'is_optional': False}
                    )
                    if ss_new:
                        created_ss += 1
                except Exception:
                    # ignore per-student failures
                    pass
                try:
                    g, g_new = Grade.objects.get_or_create(
                        student=student,
                        class_subject=cs,
                        term=selected_term_obj,
                        sequence=selected_sequence_obj,
                        defaults={
                            'value': 0.0,
                            'status': 'draft',
                            'created_by': request.user,
                            'updated_by': request.user
                        }
                    )
                    if g_new:
                        created_grades += 1
                except Exception:
                    pass

            grades = Grade.objects.filter(
                class_subject=cs,
                term=selected_term_id,
                sequence=selected_sequence_id
            ).select_related('student').order_by('student__last_name', 'student__first_name')
        else:
            # schoolyear selected but missing term/sequence -> do not display list
            grades = None
    else:
        # no schoolyear selected: do not show list
        grades = None
    auto_open_edit = request.GET.get('open_edit') == '1'
    # permission flag for template (teacher assigned to subject can manage)
    can_manage = user_can_manage_cs(request.user, cs)
    # admin/secretary check for edit rights on validated/locked grades
    is_admin_sec = is_admin_or_secretary(request.user)
    return render(request, 'Bull/classsubject_students.html', {
        'cs': cs,
        'students': students,
        'schoolyears': schoolyears,
        'terms': terms,
        'sequences': sequences,
        'selected_schoolyear': selected_schoolyear_id,
        'selected_term': selected_term_id,
        'selected_sequence': selected_sequence_id,
        'selected_schoolyear_obj': selected_schoolyear_obj,
        'selected_term_obj': selected_term_obj,
        'selected_sequence_obj': selected_sequence_obj,
        'grades': grades,
    'auto_open_edit': auto_open_edit,
    'can_manage': can_manage,
    'is_admin_or_secretary': is_admin_sec,
        'message': request.GET.get('success'),
        'error_message': request.GET.get('error'),
    })

@login_required
@user_passes_test(is_admin_or_secretary)
def generate_grades_view(request, cs_id):
    from Bull.models import ClassSubject, Grade, Sequence
    cs = ClassSubject.objects.select_related('classroom', 'subject').get(id=cs_id)
    students = cs.classroom.students.all().order_by('last_name', 'first_name')
    term_id = request.GET.get('term')
    sequence_id = request.GET.get('sequence')
    if not (term_id and sequence_id):
        url = reverse('classsubject_students', args=[cs_id]) + '?error=Veuillez sélectionner le trimestre et la séquence.'
        return redirect(url)
    from Bull.models import Term
    term = Term.objects.get(id=term_id)
    sequence = Sequence.objects.get(id=sequence_id)
    created_count = 0
    updated_count = 0
    ss_created = 0
    error_details = []
    from Bull.models import StudentSubject
    for student in students:
        # ensure StudentSubject exists
        try:
            ss, ss_new = StudentSubject.objects.get_or_create(
                student=student,
                subject=cs.subject,
                defaults={'is_optional': False}
            )
            if ss_new:
                ss_created += 1
        except Exception as e:
            error_details.append(f"StudentSubject {student.id}: {str(e)}")
            # continue to try creating grades even if student-subject linkage fails
        try:
            obj, created = Grade.objects.get_or_create(
                student=student,
                class_subject=cs,
                term=term,
                sequence=sequence,
                defaults={
                    'value': 0.0,
                    'status': 'draft',
                    'created_by': request.user,
                    'updated_by': request.user
                }
            )
            if created:
                created_count += 1
            else:
                obj.term = term
                obj.sequence = sequence
                obj.updated_by = request.user
                obj.save()
                updated_count += 1
        except Exception as e:
            error_details.append(f"Grade {student.id}: {str(e)}")
    # Redirige avec tous les paramètres pour afficher la liste générée
    params = f'?term={term_id}&sequence={sequence_id}'
    if request.GET.get('schoolyear'):
        params += f'&schoolyear={request.GET.get("schoolyear")}'
    # If caller requested generate+fill, instruct the classsubject page to open edit mode
    if request.GET.get('action') == 'generate_and_fill' or request.GET.get('action') == 'generate_and_edit':
        params += '&open_edit=1'
    if error_details:
        params += f'&error=Erreur création notes: {'; '.join(error_details)}'
    else:
        params += f'&success={created_count} notes créées, {updated_count} notes mises à jour.'
    url = reverse('classsubject_students', args=[cs_id]) + params
    return redirect(url)

@login_required
@user_passes_test(is_admin_or_secretary)
def save_grades_view(request, cs_id):
    from Bull.models import Grade
    if request.method != 'POST':
        return redirect('classsubject_students', cs_id=cs_id)
    cs = ClassSubject.objects.get(id=cs_id)
    sequence_id = request.POST.get('sequence')
    term_id = request.POST.get('term')
    saved = 0
    skipped_locked = 0
    skipped_validated = 0
    errors = []
    # expect inputs named grade_<id>
    for key, val in request.POST.items():
        if not key.startswith('grade_'):
            continue
        try:
            gid = int(key.split('_', 1)[1])
        except Exception:
            continue
        try:
            grade = Grade.objects.get(id=gid, class_subject=cs)
        except Grade.DoesNotExist:
            errors.append(f"Grade {gid} introuvable")
            continue
        try:
            # normalize empty to 0
            if val is None or val == '':
                num = 0.0
            else:
                num = float(val)
            # skip if grade locked
            if grade.status == 'locked':
                skipped_locked += 1
                continue
            # skip if grade validated and user is not admin/secretary
            if grade.status == 'validated' and not is_admin_or_secretary(request.user):
                skipped_validated += 1
                continue
            # basic validation
            if num < 0 or num > 20:
                errors.append(f"Valeur invalide pour {grade.student}: {num}")
                continue
            grade.value = num
            grade.updated_by = request.user
            # keep status as draft; you can change to validated if desired
            grade.save()
            saved += 1
        except Exception as e:
            errors.append(f"Erreur sauvegarde grade {gid}: {str(e)}")
    params = ''
    if term_id:
        params += f'?term={term_id}&sequence={sequence_id}'
    if errors:
        params += ('&' if params else '?') + 'error=' + '%20'.join(errors)
    else:
        msg = f'{saved} notes enregistrées'
        if skipped_locked:
            msg += f', {skipped_locked} notes verrouillées ignorées'
        if skipped_validated:
            msg += f', {skipped_validated} notes validées (modifiables seulement par admin/secrétariat) ignorées'
        params += ('&' if params else '?') + f'success={msg}'
    return redirect(reverse('classsubject_students', args=[cs_id]) + params)

@login_required
@user_passes_test(is_admin_or_secretary)
def validate_grades_view(request, cs_id):
    from Bull.models import Grade
    cs = ClassSubject.objects.get(id=cs_id)
    term_id = request.GET.get('term')
    sequence_id = request.GET.get('sequence')
    if not (term_id and sequence_id):
        return redirect(reverse('classsubject_students', args=[cs_id]) + '?error=Sélectionnez trimestre et séquence pour valider')
    from Bull.models import Grade
    grades = Grade.objects.filter(class_subject=cs, term_id=term_id, sequence_id=sequence_id)
    # prevent validation if any student has value <= 0 or null -> require strict > 0
    from django.db import models as djmodels
    bad_exists = grades.filter(djmodels.Q(value__lte=0) | djmodels.Q(value__isnull=True)).exists()
    if bad_exists:
        return redirect(reverse('classsubject_students', args=[cs_id]) + f'?error=Impossible de valider: un ou plusieurs élèves ont une note nulle ou ≤ 0')
    # mark as validated (do not lock) and record who validated
    for grade in grades:
        grade.status = 'validated'
        grade.updated_by = request.user
        grade.validated_by = request.user
        grade.save()
    updated = grades.count()
    return redirect(reverse('classsubject_students', args=[cs_id]) + f'?success={updated} notes validées')

@login_required
@user_passes_test(is_admin_or_secretary)
def download_grades_pdf(request, cs_id):
    from Bull.models import Grade, Sequence, Term, SchoolYear
    from django.http import HttpResponse
    # require term and sequence like the listing
    term_id = request.GET.get('term')
    sequence_id = request.GET.get('sequence')
    schoolyear_id = request.GET.get('schoolyear')
    if not (term_id and sequence_id and schoolyear_id):
        return redirect(reverse('classsubject_students', args=[cs_id]) + f'?error=Veuillez sélectionner année/trimestre/séquence pour exporter')
    cs = ClassSubject.objects.select_related('classroom', 'subject').get(id=cs_id)
    term = Term.objects.get(id=term_id)
    sequence = Sequence.objects.get(id=sequence_id)
    schoolyear = SchoolYear.objects.get(id=schoolyear_id)
    grades = Grade.objects.filter(class_subject=cs, term=term, sequence=sequence).select_related('student').order_by('student__last_name', 'student__first_name')

    # generate PDF using reportlab
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
    except Exception as e:
        # reportlab not installed
        return redirect(reverse('classsubject_students', args=[cs_id]) + f'?error=Package reportlab requis pour exporter en PDF: {e}')

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elems = []
    title = Paragraph(f"Notes - {cs.subject.name} - {cs.classroom.name}", styles['Title'])
    info = Paragraph(f"Année: {schoolyear.name} — Trimestre: {term.name} — Séquence: {sequence.name}", styles['Normal'])
    elems.append(title)
    elems.append(Spacer(1, 12))
    elems.append(info)
    elems.append(Spacer(1, 12))

    data = [["Élève", "Note", "Coefficient", "Statut"]]
    for g in grades:
        name = f"{g.student.last_name} {g.student.first_name}"
        value = g.value if g.value is not None else ""
        coef = getattr(g.class_subject, 'coefficient', '')
        status = g.status
        data.append([name, str(value), str(coef), status])

    table = Table(data, colWidths=[200, 60, 80, 80])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
    ]))
    elems.append(table)
    doc.build(elems)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    filename = f"notes_{cs.classroom.name}_{cs.subject.name}_{sequence.name}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response




# ---------------------------------------------------
#    GESTION DES bulletins 
# ---------------------------------------------------


@login_required
def bulletin_view(request):
    # Filtrage
    schoolyears = SchoolYear.objects.all()
    terms = Term.objects.all()
    sequences = Sequence.objects.all()
    classrooms = Classroom.objects.all()

    # Statistiques par salle pour affichage global
    classroom_stats = []
    selected_schoolyear = request.GET.get('schoolyear')
    selected_sequence = request.GET.get('sequence')
    for classroom in classrooms:
        total_students = classroom.students.count()
        # Si pas de sélection, on affiche le nombre de bulletins existants pour la salle (toutes années et séquences)
        if selected_sequence and selected_schoolyear:
            bulletins_count = Bulletin.objects.filter(classroom=classroom, sequence_id=selected_sequence, sequence__term__school_year_id=selected_schoolyear).count()
        else:
            bulletins_count = Bulletin.objects.filter(classroom=classroom).count()
        classroom_stats.append({
            'name': classroom.name,
            'id': classroom.id,
            'total_students': total_students,
            'bulletins_count': bulletins_count
        })

    selected_schoolyear = request.GET.get('schoolyear')
    selected_sequence = request.GET.get('sequence')
    selected_classroom = request.GET.get('classroom')

    # Si aucune année n'est sélectionnée, prendre l'année active
    if not selected_schoolyear:
        active_sy = SchoolYear.objects.filter(is_active=True).first()
        if active_sy:
            selected_schoolyear = str(active_sy.id)

    selected_schoolyear_obj = SchoolYear.objects.filter(id=selected_schoolyear).first() if selected_schoolyear else None
    selected_sequence_obj = Sequence.objects.filter(id=selected_sequence).first() if selected_sequence else None
    selected_classroom_obj = Classroom.objects.filter(id=selected_classroom).first() if selected_classroom else None

    students = Student.objects.filter(classroom_id=selected_classroom).order_by('last_name', 'first_name') if selected_classroom else []
    subjects = Subject.objects.filter(classsubject__classroom_id=selected_classroom).distinct() if selected_classroom else Subject.objects.none()

    grades = []
    grades_matrix = {}
    if selected_classroom and selected_sequence:
        term_obj = selected_sequence_obj.term if selected_sequence_obj else None
        grades = Grade.objects.filter(
            class_subject__classroom_id=selected_classroom,
            sequence_id=selected_sequence,
            term=term_obj
        ).select_related('student', 'class_subject')

        # Préparer un dictionnaire: {student_id: {subject_id: grade}}
        for student in students:
            grades_matrix[student.id] = {}
        for g in grades:
            grades_matrix[g.student_id][g.class_subject.subject_id] = g

    # Construire une structure directement exploitable dans le template
    student_grades = []
    for student in students:
        row = {"student": student, "grades": []}
        for subject in subjects:
            grade = grades_matrix.get(student.id, {}).get(subject.id)
            row["grades"].append(grade)
        student_grades.append(row)

    bulletins = Bulletin.objects.filter(
        sequence_id=selected_sequence,
        student__classroom_id=selected_classroom,
        sequence__term__school_year_id=selected_schoolyear
    ) if selected_sequence and selected_classroom and selected_schoolyear else Bulletin.objects.none()

    can_calculate = request.user.role in ['admin', 'secretary'] and students and subjects and grades and all(
        bulletin_tags.get_all_ok(grades, s.id) for s in students
    )
    can_export = bulletins.exists()

    context = {
        'schoolyears': schoolyears,
        'terms': terms,
        'sequences': sequences,
        'classrooms': classrooms,
        'selected_schoolyear': selected_schoolyear,
        'selected_sequence': selected_sequence,
        'selected_classroom': selected_classroom,
        'selected_schoolyear_obj': selected_schoolyear_obj,
        'selected_sequence_obj': selected_sequence_obj,
        'selected_classroom_obj': selected_classroom_obj,
        'students': students,
        'subjects': subjects,
        'grades': grades,
        'student_grades': student_grades,  
        'bulletins': bulletins,
        'can_calculate': can_calculate,
        'can_export': can_export,
        'classroom_stats': classroom_stats,
    }
    return render(request, 'Bull/bulletins.html', context)

@login_required
@user_passes_test(lambda u: u.role in ['admin', 'secretary'])
def calculate_bulletins(request):
    if request.method != 'POST':
        return redirect('bulletins')
    sequence_id = request.POST.get('sequence')
    classroom_id = request.POST.get('classroom')
    sequence = Sequence.objects.get(id=sequence_id)
    classroom = Classroom.objects.get(id=classroom_id)
    students = Student.objects.filter(classroom=classroom).order_by('last_name', 'first_name')
    subjects = Subject.objects.filter(classsubject__classroom=classroom).distinct()
    class_subjects = ClassSubject.objects.filter(classroom=classroom)
    grades = Grade.objects.filter(class_subject__in=class_subjects, sequence=sequence)
    # Vérification : tous les élèves doivent avoir une note valide à chaque matière
    missing = False
    for student in students:
        for subject in subjects:
            cs = class_subjects.filter(subject=subject).first()
            grade = grades.filter(student=student, class_subject=cs).first()
            if not grade or grade.status != 'validated':
                missing = True
                break
        if missing:
            break
    if missing:
        msg = "Impossible de générer les bulletins : certaines notes sont manquantes ou non validées."
        from django.contrib import messages
        messages.error(request, msg)
        return redirect(f"{reverse('bulletins')}?sequence={sequence_id}&classroom={classroom_id}")
    # Récupération du canevas bulletin (entête et pied de page) au format Word
    from Bull.models import BulletinTemplate
    canevas = BulletinTemplate.objects.filter(active=True).first()
    from docx import Document
    entete_text = ""
    pied_text = ""
    # Extraction texte header_docx
    if canevas and getattr(canevas, 'header_docx', None):
        try:
            doc = Document(canevas.header_docx.path)
            entete_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            entete_text = ""
    # Extraction texte footer_docx
    if canevas and getattr(canevas, 'footer_docx', None):
        try:
            doc = Document(canevas.footer_docx.path)
            pied_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            pied_text = ""

    # Calcul des moyennes et stats
    results = []
    averages = []
    recap_bulletins = []
    for student in students:
        total = 0
        total_coef = 0
        student_grades = [g for g in grades if g.student_id == student.id]
        recap_notes = []
        for subject in subjects:
            grade = next((g for g in student_grades if g.class_subject.subject_id == subject.id), None)
            coef = ClassSubject.objects.get(classroom=classroom, subject=subject).coefficient if ClassSubject.objects.filter(classroom=classroom, subject=subject).exists() else 1
            note = grade.value if grade and grade.value is not None else 0
            total += note * coef
            total_coef += coef
            # Rang par matière
            grades_matiere = [g for g in grades if g.class_subject.subject_id == subject.id]
            notes_matiere = sorted([g.value if g.value is not None else 0 for g in grades_matiere], reverse=True)
            rang_matiere = notes_matiere.index(note) + 1 if note in notes_matiere else '-'
            recap_notes.append({
                'matiere': subject.name,
                'note': note,
                'coef': coef,
                'som_coef': total_coef,
                'rang_matiere': rang_matiere,
                'rang_general': None, # sera rempli après
            })
        avg = round(total / total_coef, 2) if total_coef > 0 else 0
        averages.append(avg)
        results.append({'student': student, 'average': avg})
        recap_bulletins.append({
            'student': student,
            'recap_notes': recap_notes,
            'average': avg,
            'total_coef': total_coef,
        })
    # Stats simples
    moyenne_generale = round(sum(averages) / len(averages), 2) if averages else 0
    moyenne_min = min(averages) if averages else 0
    moyenne_max = max(averages) if averages else 0
    nb_echec = len([a for a in averages if a < 10])
    # Calcul du rang général
    results.sort(key=lambda x: x['average'], reverse=True)
    for idx, res in enumerate(results, 1):
        res['rank'] = idx
    # Ajout du rang général dans recap_bulletins
    for recap in recap_bulletins:
        student = recap['student']
        rank = next((r['rank'] for r in results if r['student'] == student), None)
        for note in recap['recap_notes']:
            note['rang_general'] = rank

    # Génération des bulletins PDF avec ReportLab (pur Python)
    pdf_dir = os.path.join(settings.MEDIA_ROOT, 'bulletins')
    os.makedirs(pdf_dir, exist_ok=True)
    for recap in recap_bulletins:
        student = recap['student']
        avg = recap['average']
        rank = next((r['rank'] for r in results if r['student'] == student), None)
        pdf_path = os.path.join(pdf_dir, f"bulletin_{student.id}_{sequence.id}.pdf")
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4
        # Header (entête Word)
        c.setFont("Helvetica-Bold", 12)
        y_header = height - 40
        for line in entete_text.split("\n"):
            c.drawString(50, y_header, line)
            y_header -= 16
        # Titre bulletin
        y_title = y_header - 20
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, y_title, f"Bulletin de {student.last_name} {student.first_name}")
        c.setFont("Helvetica", 12)
        c.drawString(100, y_title-20, f"Classe : {classroom.name} | Séquence : {sequence.name}")
        c.drawString(100, y_title-40, f"Moyenne : {avg} | Rang : {rank}")
        # Tableau de notes (centré)
        y_table = y_title-70
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_table, "Matière")
        c.drawString(200, y_table, "Note")
        c.drawString(260, y_table, "Coef")
        c.drawString(320, y_table, "Somme Coef")
        c.drawString(420, y_table, "Rang matière")
        c.drawString(520, y_table, "Rang général")
        y_table -= 20
        c.setFont("Helvetica", 12)
        for note in recap['recap_notes']:
            c.drawString(100, y_table, str(note['matiere']))
            c.drawString(200, y_table, str(note['note']))
            c.drawString(260, y_table, str(note['coef']))
            c.drawString(320, y_table, str(note['som_coef']))
            c.drawString(420, y_table, str(note['rang_matiere']))
            c.drawString(520, y_table, str(note['rang_general']))
            y_table -= 20
        # Footer (pied Word)
        c.setFont("Helvetica-Oblique", 11)
        y_footer = 40
        for line in pied_text.split("\n"):
            c.drawString(50, y_footer, line)
            y_footer += 16
        c.save()
        Bulletin.objects.update_or_create(
            student=student,
            classroom=classroom,
            sequence=sequence,
            defaults={
                'pdf_path': os.path.relpath(pdf_path, settings.MEDIA_ROOT),
                'average': avg,
                'rank': rank,
            }
        )
    # Lock all grades after generation
    grades.update(status='locked')

    # Rendu HTML du bulletin pour chaque élève (optionnel, pour consultation ou export)
    # Exemple pour le premier élève (à adapter selon besoin)
    if recap_bulletins:
        bulletin_data = recap_bulletins[0]
        return render(request, 'Bull/bulletin.html', {
            'recap_notes': bulletin_data['recap_notes'],
            'average': bulletin_data['average'],
            'total_coef': bulletin_data['total_coef'],
        })
    
# 3. Vue détail du bulletin
@login_required
def bulletin_detail_view(request, student_id, sequence_id):
    student = get_object_or_404(Student, id=student_id)
    sequence = get_object_or_404(Sequence, id=sequence_id)
    classroom = student.classroom
    # Inclure toutes les matières ayant au moins une note pour cette classe et séquence
    from Bull.models import Grade, ClassSubject, Subject, Bulletin
    subjects = Subject.objects.filter(
        id__in=Grade.objects.filter(
            class_subject__classroom=classroom,
            sequence=sequence
        ).values_list('class_subject__subject_id', flat=True)
    ).distinct()
    class_subjects = ClassSubject.objects.filter(classroom=classroom, subject__in=subjects)
    grades = Grade.objects.filter(student=student, class_subject__in=class_subjects, sequence=sequence)
    bulletin = Bulletin.objects.filter(student=student, sequence=sequence).first()
    # Calcul moyenne et rang
    avg = bulletin.average if bulletin else None
    rank = bulletin.rank if bulletin else None
    pdf_path = bulletin.pdf_path if bulletin else None
    # Calcul detailed_row pour le template
    notes = []
    notes_x_coef = []
    somme = 0
    total_coef = 0
    rangs_matiere = {}
    # Rang par matière
    for subject in subjects:
        cs = class_subjects.filter(subject=subject).first()
        grades_matiere = Grade.objects.filter(class_subject=cs, sequence=sequence)
        notes_matiere = []
        for g in grades_matiere:
            note = g.value if g.value is not None else 0
            notes_matiere.append({'student_id': g.student_id, 'note': note})
        notes_matiere.sort(key=lambda x: x['note'], reverse=True)
        for idx, entry in enumerate(notes_matiere, 1):
            rangs_matiere[(subject.id, entry['student_id'])] = idx
    for subject in subjects:
        cs = class_subjects.filter(subject=subject).first()
        grade = grades.filter(class_subject=cs).first()
        note = grade.value if grade and grade.value is not None else 0
        coef = cs.coefficient if cs else 1
        note_x_coef = round(note * coef, 2)
        rang_matiere = rangs_matiere.get((subject.id, student.id), '-')
        notes.append({'subject': subject.name, 'note': note, 'coef': coef, 'note_x_coef': note_x_coef, 'rang_matiere': rang_matiere})
        notes_x_coef.append(note_x_coef)
        somme += note_x_coef
        total_coef += coef
    moyenne = round(somme / total_coef, 2) if total_coef > 0 else 0
    somme_notes = round(sum([n['note'] for n in notes]), 2)
    detailed_row = {
        'nom': f"{student.last_name} {student.first_name}",
        'notes': notes,
        'notes_x_coef': notes_x_coef,
        'somme': somme,
        'somme_notes': somme_notes,
        'total_coef': total_coef,
        'moyenne': moyenne,
        'rang': rank,
    }
    return render(request, 'Bull/bulletin_detail.html', {
        'student': student,
        'sequence': sequence,
        'classroom': classroom,
        'subjects': subjects,
        'grades': grades,
        'avg': avg,
        'rank': rank,
        'pdf_path': pdf_path,
        'detailed_row': detailed_row,
    })
    # Génération Excel
    import openpyxl
    excel_path = os.path.join(pdf_dir, f"bulletins_classe_{classroom.id}_seq_{sequence.id}.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bulletins"
    ws.append(["Élève", "Moyenne", "Rang"])
    for res in results:
        student = res['student']
        ws.append([f"{student.last_name} {student.first_name}", res['average'], res['rank']])
    wb.save(excel_path)
    # Stocke les stats en session pour la page suivante
    request.session['bulletin_stats'] = {
        'moyenne_generale': moyenne_generale,
        'moyenne_min': moyenne_min,
        'moyenne_max': moyenne_max,
        'nb_echec': nb_echec,
        'results': [
            {'nom': f"{r['student'].last_name} {r['student'].first_name}", 'moyenne': r['average'], 'rang': r['rank']} for r in results
        ],
        'excel_path': excel_path,
    }
    return redirect('bulletin_stats')

@login_required
def bulletin_stats(request):
    # Récupère les paramètres de classe et séquence
    classroom_id = request.GET.get('classroom')
    sequence_id = request.GET.get('sequence')
    stats = {}
    if classroom_id and sequence_id:
        from Bull.models import Bulletin, Student, Classroom, Sequence, Subject, ClassSubject, Grade
        classroom = Classroom.objects.filter(id=classroom_id).first()
        sequence = Sequence.objects.filter(id=sequence_id).first()
        bulletins = Bulletin.objects.filter(classroom_id=classroom_id, sequence_id=sequence_id)
        students = Student.objects.filter(classroom_id=classroom_id).order_by('last_name', 'first_name')
        subjects = Subject.objects.filter(
            id__in=Grade.objects.filter(
                class_subject__classroom_id=classroom_id,
                sequence_id=sequence_id
            ).values_list('class_subject__subject_id', flat=True)
        ).distinct()
        class_subjects = ClassSubject.objects.filter(classroom_id=classroom_id, subject__in=subjects)
        results = []
        averages = []
        detailed_rows = []
        # Calcul du rang par matière
        rangs_matiere = {}
        for subject in subjects:
            # Pour chaque matière, on récupère toutes les notes des élèves
            cs = class_subjects.filter(subject=subject).first()
            grades_matiere = Grade.objects.filter(class_subject=cs, sequence_id=sequence_id)
            notes_matiere = []
            for student in students:
                grade = grades_matiere.filter(student=student).first()
                note = grade.value if grade and grade.value is not None else 0
                notes_matiere.append({'student_id': student.id, 'note': note})
            # Classement décroissant
            notes_matiere.sort(key=lambda x: x['note'], reverse=True)
            for idx, entry in enumerate(notes_matiere, 1):
                rangs_matiere[(subject.id, entry['student_id'])] = idx

        for student in students:
            b = bulletins.filter(student=student).first()
            grades = Grade.objects.filter(student=student, class_subject__in=class_subjects, sequence_id=sequence_id)
            notes = []
            notes_x_coef = []
            somme = 0
            total_coef = 0
            for subject in subjects:
                cs = class_subjects.filter(subject=subject).first()
                grade = grades.filter(class_subject=cs).first()
                note = grade.value if grade and grade.value is not None else 0
                coef = cs.coefficient if cs else 1
                rang_matiere = rangs_matiere.get((subject.id, student.id), '-')
                notes.append({'subject': subject.name, 'note': note, 'coef': coef, 'rang_matiere': rang_matiere})
                notes_x_coef.append(note * coef)
                somme += note * coef
                total_coef += coef
            moyenne = round(somme / total_coef, 2) if total_coef > 0 else 0
            avg = b.average if b and b.average is not None else moyenne
            averages.append(avg)
            rang = b.rank if b else None
            results.append({'student': student, 'average': avg, 'rank': rang})
            detailed_rows.append({
                'nom': f"{student.last_name} {student.first_name}",
                'notes': notes,
                'notes_x_coef': notes_x_coef,
                'somme': somme,
                'total_coef': total_coef,
                'moyenne': moyenne,
                'rang': rang,
            })
        # Stats simples
        moyenne_generale = round(sum(averages) / len(averages), 2) if averages else 0
        moyenne_min = min(averages) if averages else 0
        moyenne_max = max(averages) if averages else 0
        nb_echec = len([a for a in averages if a < 10])
        # Calcul du rang
        results.sort(key=lambda x: x['average'], reverse=True)
        for idx, res in enumerate(results, 1):
            res['rank'] = idx
        # Format pour le template
        stats = {
            'moyenne_generale': moyenne_generale,
            'moyenne_min': moyenne_min,
            'moyenne_max': moyenne_max,
            'nb_echec': nb_echec,
            'results': [
                {'nom': f"{r['student'].last_name} {r['student'].first_name}", 'moyenne': r['average'], 'rang': r['rank']} for r in results
            ],
            'detailed_rows': detailed_rows,
            'subjects': [s.name for s in subjects],
            'excel_path': None,
        }
    return render(request, 'Bull/bulletin_stats.html', {'stats': stats})

@login_required
def export_bulletins_pdf(request):
    sequence_id = request.GET.get('sequence')
    classroom_id = request.GET.get('classroom')
    bulletins = Bulletin.objects.filter(classroom_id=classroom_id, sequence_id=sequence_id)
    import zipfile
    from io import BytesIO
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for b in bulletins:
            if b.pdf_path and hasattr(b.pdf_path, 'path') and os.path.exists(b.pdf_path.path):
                filename = f"{b.student.last_name}_{b.student.first_name}_bulletin_{sequence_id}.pdf"
                zip_file.write(b.pdf_path.path, arcname=filename)
    zip_buffer.seek(0)
    return FileResponse(zip_buffer, as_attachment=True, filename=f"bulletins_classe_{classroom_id}_seq_{sequence_id}.zip")

@login_required
def export_bulletins_excel(request):
    import openpyxl
    sequence_id = request.GET.get('sequence')
    classroom_id = request.GET.get('classroom')
    from Bull.models import Bulletin, Student, Subject, ClassSubject, Grade, Sequence, Classroom
    bulletins = Bulletin.objects.filter(classroom_id=classroom_id, sequence_id=sequence_id)
    students = Student.objects.filter(classroom_id=classroom_id).order_by('last_name', 'first_name')
    classroom = Classroom.objects.filter(id=classroom_id).first()
    sequence = Sequence.objects.filter(id=sequence_id).first()
    subjects = Subject.objects.filter(classsubject__classroom_id=classroom_id).distinct()
    class_subjects = ClassSubject.objects.filter(classroom_id=classroom_id, subject__in=subjects)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bulletins"
    # Titre
    ws.append([f"Classe : {classroom.name if classroom else classroom_id}", f"Séquence : {sequence.name if sequence else sequence_id}"])
    # En-tête
    header = ["Nom"]
    for subject in subjects:
        header.append(subject.name)
    for subject in subjects:
        header.append(f"{subject.name} x coef")
    header += ["Somme", "Total coef", "Moyenne", "Rang"]
    ws.append(header)
    # Lignes élèves
    for student in students:
        b = bulletins.filter(student=student).first()
        grades = Grade.objects.filter(student=student, class_subject__in=class_subjects, sequence_id=sequence_id)
        row = [f"{student.last_name} {student.first_name}"]
        notes = []
        notes_x_coef = []
        somme = 0
        total_coef = 0
        for subject in subjects:
            cs = class_subjects.filter(subject=subject).first()
            grade = grades.filter(class_subject=cs).first()
            note = grade.value if grade and grade.value is not None else 0
            coef = cs.coefficient if cs else 1
            notes.append(note)
            notes_x_coef.append(note * coef)
            somme += note * coef
            total_coef += coef
        row += notes
        row += notes_x_coef
        row += [somme, total_coef]
        moyenne = round(somme / total_coef, 2) if total_coef > 0 else 0
        row.append(moyenne)
        rang = b.rank if b else "-"
        row.append(rang)
        ws.append(row)
    excel_path = os.path.join(settings.MEDIA_ROOT, f"bulletins/bulletins_classe_{classroom_id}_seq_{sequence_id}.xlsx")
    wb.save(excel_path)
    return FileResponse(open(excel_path, 'rb'), as_attachment=True, filename=os.path.basename(excel_path))



@login_required
def download_bulletin_pdf(request, student_id, sequence_id):
    bulletin = Bulletin.objects.filter(student_id=student_id, sequence_id=sequence_id).first()
    if not bulletin or not bulletin.pdf_path:
        raise Http404("Bulletin PDF introuvable.")
    pdf_abspath = bulletin.pdf_path.path
    if not os.path.exists(pdf_abspath):
        raise Http404("Fichier PDF introuvable.")
    return FileResponse(open(pdf_abspath, 'rb'), as_attachment=False, filename=os.path.basename(pdf_abspath))





# Vue de consultation des bulletins consolidés
@login_required
@user_passes_test(lambda u: u.role in ['admin', 'secretary'])
def consolidated_bulletins(request):
    classroom_id = request.GET.get('classroom')
    schoolyear_id = request.GET.get('schoolyear')
    classroom = Classroom.objects.filter(id=classroom_id).first()
    schoolyear = SchoolYear.objects.filter(id=schoolyear_id).first()
    students = Student.objects.filter(classroom=classroom).order_by('last_name', 'first_name') if classroom else []
    terms = schoolyear.terms.all() if schoolyear else []
    sequences = Sequence.objects.filter(term__in=terms) if terms else []
    # Récupère tous les bulletins (séquence, trimestre, annuel) pour chaque élève
    bulletins_by_student = {}
    for student in students:
        bulletins = Bulletin.objects.filter(student=student).order_by('sequence__order')
        bulletins_by_student[student.id] = bulletins
    context = {
        'classroom': classroom,
        'schoolyear': schoolyear,
        'students': students,
        'terms': terms,
        'sequences': sequences,
        'bulletins_by_student': bulletins_by_student,
    }
    return render(request, 'Bull/consolidated_bulletins.html', context)
# Génération des bulletins trimestriels
@login_required
@user_passes_test(lambda u: u.role in ['admin', 'secretary'])
def generate_bulletins_trimester(request):
    if request.method != 'POST':
        return redirect('bulletins')
    classroom_id = request.POST.get('classroom')
    schoolyear_id = request.POST.get('schoolyear')
    trimester_id = request.POST.get('trimester')
    classroom = Classroom.objects.get(id=classroom_id)
    term = Term.objects.get(id=trimester_id)
    schoolyear = SchoolYear.objects.get(id=schoolyear_id)
    sequences = term.sequences.all()
    students = Student.objects.filter(classroom=classroom).order_by('last_name', 'first_name')
    # Vérification : tous les élèves ont un bulletin pour chaque séquence du trimestre
    missing = []
    for student in students:
        for seq in sequences:
            if not Bulletin.objects.filter(student=student, sequence=seq).exists():
                missing.append(f"{student.last_name} {student.first_name} - {seq.name}")
    if missing:
        messages.error(request, "Impossible de générer le bulletin trimestriel : certains bulletins de séquence sont manquants.")
        return redirect(f"{reverse('bulletins')}?classroom={classroom_id}&schoolyear={schoolyear_id}")
    # Calcul des moyennes, rangs, remarques, génération PDF individuel
    for student in students:
        # Récupère les bulletins de séquence
        seq_bulletins = Bulletin.objects.filter(student=student, sequence__in=sequences)
        moyennes_seq = [b.average for b in seq_bulletins if b.average is not None]
        moyenne_trim = round(sum(moyennes_seq) / len(moyennes_seq), 2) if moyennes_seq else 0
        # Rang du trimestre
        all_moyennes = []
        for s in students:
            seq_bulletins_s = Bulletin.objects.filter(student=s, sequence__in=sequences)
            moyennes_seq_s = [b.average for b in seq_bulletins_s if b.average is not None]
            m = round(sum(moyennes_seq_s) / len(moyennes_seq_s), 2) if moyennes_seq_s else 0
            all_moyennes.append((s.id, m))
        all_moyennes.sort(key=lambda x: x[1], reverse=True)
        rank = next((i+1 for i, (sid, m) in enumerate(all_moyennes) if sid == student.id), None)
        # Appréciation
        appreciation = "Excellent" if moyenne_trim >= 16 else "Très bien" if moyenne_trim >= 14 else "Bien" if moyenne_trim >= 12 else "Passable" if moyenne_trim >= 10 else "Insuffisant"
        # Génération PDF individuel
        pdf_dir = os.path.join(settings.MEDIA_ROOT, 'bulletins')
        os.makedirs(pdf_dir, exist_ok=True)
        pdf_path = os.path.join(pdf_dir, f"bulletin_trim_{student.id}_{term.id}.pdf")
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.drawString(100, 800, f"Bulletin Trimestriel de {student.last_name} {student.first_name}")
        c.drawString(100, 780, f"Classe : {classroom.name} | Trimestre : {term.name}")
        c.drawString(100, 760, f"Moyenne trimestre : {moyenne_trim}")
        c.drawString(100, 740, f"Rang trimestre : {rank}")
        c.drawString(100, 720, f"Appréciation : {appreciation}")
        y = 700
        for b in seq_bulletins:
            c.drawString(100, y, f"Séquence {b.sequence.name} : Moyenne {b.average} | Rang {b.rank}")
            y -= 20
        c.save()
        # Enregistre le bulletin trimestriel (marqué)
        Bulletin.objects.create(
            student=student,
            classroom=classroom,
            sequence=sequences.first(), # pour référence, mais on peut ajouter un champ spécifique
            pdf_path=os.path.relpath(pdf_path, settings.MEDIA_ROOT),
            average=moyenne_trim,
            rank=rank,
            comment=appreciation,
            # Marquage type bulletin
            is_trimester=True
        )
    # Redirection vers la page de consultation des bulletins consolidés
    return redirect(f"{reverse('consolidated_bulletins')}?classroom={classroom_id}&schoolyear={schoolyear_id}")

# Génération des bulletins annuels
@login_required
@user_passes_test(lambda u: u.role in ['admin', 'secretary'])
def generate_bulletins_annual(request):
    if request.method != 'POST':
        return redirect('bulletins')
    classroom_id = request.POST.get('classroom')
    schoolyear_id = request.POST.get('schoolyear')
    classroom = Classroom.objects.get(id=classroom_id)
    schoolyear = SchoolYear.objects.get(id=schoolyear_id)
    terms = schoolyear.terms.all()
    sequences = Sequence.objects.filter(term__in=terms)
    students = Student.objects.filter(classroom=classroom).order_by('last_name', 'first_name')
    # Vérification : tous les élèves ont une note pour chaque séquence de l'année
    missing = []
    for student in students:
        for seq in sequences:
            if not Bulletin.objects.filter(student=student, sequence=seq).exists():
                missing.append(f"{student.last_name} {student.first_name} - {seq.name}")
    if missing:
        messages.error(request, "Impossible de générer le bulletin annuel : certains bulletins de séquence sont manquants.")
        return redirect(f"{reverse('bulletins')}?classroom={classroom_id}&schoolyear={schoolyear_id}")
    # Calcul des moyennes, rangs, remarques, génération PDF individuel
    for student in students:
        seq_bulletins = Bulletin.objects.filter(student=student, sequence__in=sequences)
        moyennes_seq = [b.average for b in seq_bulletins if b.average is not None]
        moyenne_annuelle = round(sum(moyennes_seq) / len(moyennes_seq), 2) if moyennes_seq else 0
        # Rang annuel
        all_moyennes = []
        for s in students:
            seq_bulletins_s = Bulletin.objects.filter(student=s, sequence__in=sequences)
            moyennes_seq_s = [b.average for b in seq_bulletins_s if b.average is not None]
            m = round(sum(moyennes_seq_s) / len(moyennes_seq_s), 2) if moyennes_seq_s else 0
            all_moyennes.append((s.id, m))
        all_moyennes.sort(key=lambda x: x[1], reverse=True)
        rank = next((i+1 for i, (sid, m) in enumerate(all_moyennes) if sid == student.id), None)
        appreciation = "Excellent" if moyenne_annuelle >= 16 else "Très bien" if moyenne_annuelle >= 14 else "Bien" if moyenne_annuelle >= 12 else "Passable" if moyenne_annuelle >= 10 else "Insuffisant"
        pdf_dir = os.path.join(settings.MEDIA_ROOT, 'bulletins')
        os.makedirs(pdf_dir, exist_ok=True)
        pdf_path = os.path.join(pdf_dir, f"bulletin_annuel_{student.id}_{schoolyear.id}.pdf")
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.drawString(100, 800, f"Bulletin Annuel de {student.last_name} {student.first_name}")
        c.drawString(100, 780, f"Classe : {classroom.name} | Année : {schoolyear.name}")
        c.drawString(100, 760, f"Moyenne annuelle : {moyenne_annuelle}")
        c.drawString(100, 740, f"Rang annuel : {rank}")
        c.drawString(100, 720, f"Appréciation : {appreciation}")
        y = 700
        for b in seq_bulletins:
            c.drawString(100, y, f"Séquence {b.sequence.name} : Moyenne {b.average} | Rang {b.rank}")
            y -= 20
        c.save()
        Bulletin.objects.create(
            student=student,
            classroom=classroom,
            sequence=sequences.first(),
            pdf_path=os.path.relpath(pdf_path, settings.MEDIA_ROOT),
            average=moyenne_annuelle,
            rank=rank,
            comment=appreciation,
            is_annual=True
        )
    # Redirection vers la page de consultation des bulletins consolidés
    return redirect(f"{reverse('consolidated_bulletins')}?classroom={classroom_id}&schoolyear={schoolyear_id}")
# Vue pour servir le PDF du bulletin d'un élève



# Vue édition canevas HTML bulletin
@login_required
@user_passes_test(is_admin_or_secretary)
def edit_html_canvas(request):
    from Bull.models import BulletinTemplate
    template = BulletinTemplate.objects.filter(active=True).first()
    from Bull.forms import BulletinTemplateForm
    if not template:
        template = BulletinTemplate.objects.create(name="Canevas HTML", school_year=None, active=True)
    if request.method == 'POST':
        form = BulletinTemplateForm(request.POST, request.FILES, instance=template)
        if form.is_valid():
            form.save()
            return redirect('parameters')
    else:
        form = BulletinTemplateForm(instance=template)
    return render(request, 'Bull/edit_html_canvas.html', {'form': form, 'template': template})




# ------------------------------------
#   GESTIONS DES ENSEIGNANTS
# -------------------------------------

@login_required
def teachers_list_view(request):
    teachers = Teacher.objects.all()
    return render(request, 'Bull/teachers_list.html', {'teachers': teachers})

@login_required
def add_teacher_view(request):
    from Bull.forms import UserForm, TeacherForm, TeacherClassSubjectForm
    user_form = UserForm(request.POST or None)
    teacher_form = TeacherForm(request.POST or None)
    class_subject_form = TeacherClassSubjectForm(request.POST or None)
    created_teacher = None
    if request.method == 'POST':
        valid = user_form.is_valid() and teacher_form.is_valid()
        if valid:
            user = user_form.save(commit=False)
            user.role = 'teacher'
            password = user_form.cleaned_data.get('password')
            if password:
                user.set_password(password)
            user.save()
            teacher = teacher_form.save(commit=False)
            teacher.user = user
            teacher.save()
            created_teacher = teacher
        if 'add_class_subject' in request.POST and class_subject_form.is_valid() and created_teacher:
            classroom = class_subject_form.cleaned_data['classroom']
            subject = class_subject_form.cleaned_data['subject']
            from Bull.models import ClassSubject
            class_subject = ClassSubject.objects.filter(classroom=classroom, subject=subject).first()
            if class_subject:
                class_subject.teacher = created_teacher
                class_subject.save()
            return redirect('edit_teacher', teacher_id=created_teacher.id)
        if valid:
            return redirect('/teachers/')
    return render(request, 'Bull/add_teacher.html', {
        'user_form': user_form,
        'teacher_form': teacher_form,
        'class_subject_form': class_subject_form,
        'class_subjects': []
    })

@login_required
def edit_teacher_view(request, teacher_id):
    teacher = get_object_or_404(Teacher, id=teacher_id)
    user = teacher.user
    from Bull.forms import UserForm, TeacherForm, TeacherClassSubjectForm
    user_form = UserForm(request.POST or None, instance=user)
    teacher_form = TeacherForm(request.POST or None, instance=teacher)
    class_subject_form = TeacherClassSubjectForm(request.POST or None, teacher=teacher)
    if request.method == 'POST':
        valid = user_form.is_valid() and teacher_form.is_valid()
        if 'add_class_subject' in request.POST and class_subject_form.is_valid():
            classroom = class_subject_form.cleaned_data['classroom']
            subject = class_subject_form.cleaned_data['subject']
            from Bull.models import ClassSubject
            class_subject = ClassSubject.objects.filter(classroom=classroom, subject=subject).first()
            if class_subject:
                class_subject.teacher = teacher
                class_subject.save()
            return redirect('edit_teacher', teacher_id=teacher.id)
        elif valid:
            user_form.save()
            teacher_form.save()
            return redirect('/teachers/')
    return render(request, 'Bull/edit_teacher.html', {
        'user_form': user_form,
        'teacher_form': teacher_form,
        'class_subject_form': class_subject_form,
        'teacher': teacher,
        'class_subjects': teacher.classsubject_set.all()
    })

@login_required
def delete_teacher_view(request, teacher_id):
    teacher = get_object_or_404(Teacher, id=teacher_id)
    if request.method == 'POST':
        teacher.delete()
        return redirect('/teachers/')
    return render(request, 'Bull/confirm_delete_teacher.html', {'teacher': teacher})

@login_required
def teacher_detail_view(request, teacher_id):
    teacher = get_object_or_404(Teacher, id=teacher_id)
    # Récupère toutes les matières + classes associées à cet enseignant
    class_subjects = ClassSubject.objects.filter(teacher=teacher).select_related('classroom', 'subject')
    return render(request, 'Bull/teacher_detail.html', {
        'teacher': teacher,
        'class_subjects': class_subjects
    })

@login_required
@require_POST
def unlink_class_subject_teacher(request, teacher_id, classsubject_id):
    teacher = get_object_or_404(Teacher, id=teacher_id)
    class_subject = get_object_or_404(ClassSubject, id=classsubject_id, teacher=teacher)
    class_subject.teacher = None
    class_subject.save()
    return redirect('edit_teacher', teacher_id=teacher.id)
# Endpoint AJAX pour récupérer les matières liées à une classe
from django.views.decorators.http import require_GET

@require_GET
def get_subjects_for_class(request):
    classroom_id = request.GET.get('classroom_id')
    if not classroom_id:
        return JsonResponse({'subjects': []})
    subjects = Subject.objects.filter(classsubject__classroom_id=classroom_id).distinct()
    data = [{'id': s.id, 'name': s.name} for s in subjects]
    return JsonResponse({'subjects': data})
from django.http import JsonResponse
# Endpoint AJAX pour récupérer les classes liées à une matière
from Bull.models import ClassSubject, Classroom, Subject
from django.views.decorators.http import require_GET

@require_GET
def get_classes_for_subject(request):
    subject_id = request.GET.get('subject_id')
    if not subject_id:
        return JsonResponse({'classes': []})
    classes = Classroom.objects.filter(class_subjects__subject_id=subject_id).distinct()
    data = [{'id': c.id, 'name': c.name} for c in classes]
    return JsonResponse({'classes': data})


# ------------------------------------
#   GESTIONS DES ELEVES
# -------------------------------------

class StudentForm(forms.ModelForm):
    class Meta:
        model = Student
        fields = ['matricule', 'first_name', 'last_name', 'gender', 'birth_date', 'birth_place', 'classroom', 'photo']


@login_required
@user_passes_test(is_admin_or_secretary)
def students_view(request):
    classes = Classroom.objects.all()
    selected_class = request.GET.get('class')
    search = request.GET.get('search', '').strip()
    students = Student.objects.all()
    if selected_class:
        students = students.filter(classroom__id=selected_class)
    if search:
        students = students.filter(Q(matricule__icontains=search) | Q(first_name__icontains=search) | Q(last_name__icontains=search))
    return render(request, 'Bull/students.html', {
        'students': students,
        'classes': classes,
        'selected_class': selected_class,
        'search': search
    })

@login_required
@user_passes_test(is_admin_or_secretary)
def student_detail_view(request, student_id):
    student = Student.objects.get(id=student_id)
    disciplines = student.disciplines.select_related('term').all()
    return render(request, 'Bull/student_detail.html', {'student': student, 'disciplines': disciplines})

@login_required
@user_passes_test(is_admin_or_secretary)
def student_edit_view(request, student_id):
    student = Student.objects.get(id=student_id)
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES, instance=student)
        if form.is_valid():
            form.save()
            return redirect('student_detail', student_id=student.id)
    else:
        form = StudentForm(instance=student)
    return render(request, 'Bull/student_edit.html', {'form': form, 'student': student})

@login_required
def add_student_view(request):
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/students/')
    else:
        form = StudentForm()
    return render(request, 'Bull/add_student.html', {'form': form})

@login_required
def import_students_view(request):
    if request.method == 'POST':
        form = ImportStudentsForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            errors = []
            try:
                import openpyxl
                wb = openpyxl.load_workbook(excel_file)
                ws = wb.active
                for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True)):
                    try:
                        matricule, last_name, first_name, gender, birth_date, birth_place, photo, class_name, redouble = row
                        classroom = Classroom.objects.filter(name=class_name).first()
                        if not classroom:
                            errors.append(f"Ligne {i+2}: Classe '{class_name}' introuvable.")
                            continue
                        student = Student(
                            matricule=matricule,
                            last_name=last_name,
                            first_name=first_name,
                            gender=gender,
                            birth_date=birth_date,
                            birth_place=birth_place,
                            classroom=classroom,
                            repeater=(str(redouble).lower() == 'oui')
                        )
                        student.save()
                    except Exception as e:
                        errors.append(f"Ligne {i+2}: {str(e)}")
            except BadZipFile:
                errors.append("Le fichier n'est pas un vrai fichier Excel (.xlsx). Veuillez vérifier le format.")
            except Exception as e:
                errors.append(f"Erreur lors de l'import : {str(e)}")
            if errors:
                return render(request, 'Bull/students.html', {'import_errors': errors})
            return redirect('/students/')
    else:
        form = ImportStudentsForm()
    return render(request, 'Bull/students.html', {'form': form})



@login_required
def delete_student_view(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    if request.method == 'POST':
        student.delete()
        return redirect('/students/')
    # Optionnel : afficher une page de confirmation personnalisée
    return redirect(f'/students/{student_id}/')

@login_required
def delete_student_view(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    if request.method == 'POST':
        student.delete()
        return redirect('/students/')
    return render(request, 'Bull/confirm_delete_student.html', {'student': student})


@login_required
def export_students_view(request):
    classroom_id = request.GET.get('classroom')
    export_format = request.GET.get('format', 'excel')
    classroom = Classroom.objects.filter(id=classroom_id).first()
    if not classroom:
        return HttpResponse('Classe non trouvée', status=404)
    students = Student.objects.filter(classroom=classroom)

    if export_format == 'excel':
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = classroom.name
        ws.append([f"Liste des élèves de la classe {classroom.name}"])
        ws.append(['Matricule', 'Nom', 'Prénom', 'Genre', 'Date Naissance', 'Lieu Naissance', 'Redouble', 'Photo'])
        for s in students:
            ws.append([
                s.matricule,
                s.last_name,
                s.first_name,
                s.gender,
                s.birth_date.strftime('%Y-%m-%d') if s.birth_date else '',
                s.birth_place,
                'Oui' if s.repeater else 'Non',
                s.photo.url if s.photo else ''
            ])
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename=eleves_{classroom.name}.xlsx'
        wb.save(response)
        return response

    elif export_format == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename=eleves_{classroom.name}.pdf'
        p = canvas.Canvas(response, pagesize=A4)
        width, height = A4
        y = height - 50
        p.setFont('Helvetica-Bold', 16)
        p.drawString(50, y, f"Liste des élèves de la classe {classroom.name}")
        y -= 40
        p.setFont('Helvetica', 10)
        headers = ['Matricule', 'Nom', 'Prénom', 'Genre', 'Date Naissance', 'Lieu Naissance', 'Redouble']
        p.drawString(50, y, ' | '.join(headers))
        y -= 20
        for s in students:
            line = [
                s.matricule,
                s.last_name,
                s.first_name,
                s.gender,
                s.birth_date.strftime('%Y-%m-%d') if s.birth_date else '',
                s.birth_place,
                'Oui' if s.repeater else 'Non'
            ]
            p.drawString(50, y, ' | '.join(line))
            y -= 18
            if y < 50:
                p.showPage()
                y = height - 50
        p.save()
        return response

    else:
        return HttpResponse('Format non supporté', status=400)

# -----------------------------------------------
# gestion des séquences/trimestres
# -----------------------------------------------

@login_required
@user_passes_test(is_admin_or_secretary)
def add_sequence_view(request):
    from Bull.models import Term, Sequence
    terms = Term.objects.select_related('school_year').all()
    if request.method == 'POST':
        term_id = request.POST.get('term')
        name = request.POST.get('name')
        order = request.POST.get('order')
        weight = request.POST.get('weight')
        if term_id and name and order:
            term = Term.objects.get(id=term_id)
            Sequence.objects.create(term=term, name=name, order=order, weight=weight or 1)
            return redirect('dashboard')
        else:
            return render(request, 'Bull/add_sequence.html', {'terms': terms, 'error': 'Veuillez remplir tous les champs.'})
    return render(request, 'Bull/add_sequence.html', {'terms': terms})



# -----------------------------------------------------
# Fonctions de dashboard et profil
# ----------------------------------------------------


@login_required
def dashboard_view(request):
    user_role = request.user.role
    if user_role in ['admin', 'secretary']:
        nb_eleves = Student.objects.count()
        nb_enseignants = Teacher.objects.count()
        nb_classes = Classroom.objects.count()
        from Bull.models import User, SchoolYear, Term, Grade
        roles_stats = User.objects.values('role').annotate(count=Count('id'))
        school_year = SchoolYear.objects.filter(is_active=True).first()
        trimestres = Term.objects.filter(school_year=school_year) if school_year else []

        # Moyennes par classe/trimestre
        moyennes_classes = []
        if nb_classes > 0 and school_year:
            for classroom in Classroom.objects.all():
                for term in trimestres:
                    students = classroom.students.all()
                    total = 0
                    count = 0
                    for student in students:
                        avg = Grade.calculate_term_average(student, term)
                        if avg is not None:
                            total += avg
                            count += 1
                    moyenne = round(total / count, 2) if count > 0 else None
                    moyennes_classes.append({
                        'class': classroom.name,
                        'term': term.name,
                        'average': moyenne
                    })

        # Taux de réussite global et nombre d'élèves ayant réussi/échoué
        total_students = Student.objects.count()
        nb_reussite = 0
        nb_echec = 0
        if school_year:
            for student in Student.objects.all():
                avg = Grade.calculate_annual_average(student, school_year)
                if avg is not None:
                    if avg >= 10:
                        nb_reussite += 1
                    else:
                        nb_echec += 1
        taux_reussite = round((nb_reussite / total_students) * 100, 2) if total_students > 0 else None

        # Évolution des performances
        evolution = []
        for sy in SchoolYear.objects.all():
            total = 0
            count = 0
            for student in Student.objects.all():
                avg = Grade.calculate_annual_average(student, sy)
                if avg is not None:
                    total += avg
                    count += 1
            evolution.append({
                'year': sy.name,
                'average': round(total / count, 2) if count > 0 else None
            })

        # État de la BD
        from django.db import connection
        with connection.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'")
            nb_tables = cursor.fetchone()[0]

        # 1. Nombre total de bulletins générés (toutes classes, toutes séquences)
        total_bulletins = Bulletin.objects.count()

        # 2. Stats min/max par matière (notes > 0.1 uniquement)
        matiere_stats = []
        for subject in Subject.objects.all():
            grades = Grade.objects.filter(class_subject__subject=subject, value__gt=0.1)
            if grades.exists():
                min_note = grades.aggregate(models.Min('value'))['value__min']
                max_note = grades.aggregate(models.Max('value'))['value__max']
                matiere_stats.append({
                    'subject': subject.name,
                    'min': min_note,
                    'max': max_note,
                })

        # 3. Moyenne générale par classe
        moyennes_par_classe = []
        for classroom in Classroom.objects.all():
            students = classroom.students.all()
            total = 0
            count = 0
            for student in students:
                avg = Grade.objects.filter(student=student, status='validated').aggregate(models.Avg('value'))['value__avg']
                if avg is not None:
                    total += avg
                    count += 1
            moyenne = round(total / count, 2) if count > 0 else None
            moyennes_par_classe.append({'class': classroom.name, 'average': moyenne})

        # 4. Taille du dossier des bulletins générés
        bulletins_dir = os.path.join(settings.MEDIA_ROOT, 'bulletins')
        disk_usage = 0
        if os.path.exists(bulletins_dir):
            for root, dirs, files in os.walk(bulletins_dir):
                for f in files:
                    disk_usage += os.path.getsize(os.path.join(root, f))

        # 5. Idée bonus : nombre d'élèves ayant au moins un bulletin généré
        eleves_avec_bulletin = Student.objects.filter(bulletins__isnull=False).distinct().count()

        # Paramètres de l'application
        active_schoolyear = SchoolYear.objects.filter(is_active=True).first()

        context = {
            'nb_eleves': nb_eleves,
            'nb_enseignants': nb_enseignants,
            'nb_classes': nb_classes,
            'roles_stats': roles_stats,
            'school_year': school_year,
            'trimestres': trimestres,
            'moyennes_classes': moyennes_classes,
            'taux_reussite': taux_reussite,
            'nb_reussite': nb_reussite,
            'nb_echec': nb_echec,
            'evolution': evolution,
            'nb_tables': nb_tables,
            'disk_usage': disk_usage,
            'active_schoolyear': active_schoolyear,
            'dashboard_title': 'Tableau de bord Administrateur/Secrétariat',
            'total_bulletins': total_bulletins,
            'matiere_stats': matiere_stats,
            'moyennes_par_classe': moyennes_par_classe,
            'eleves_avec_bulletin': eleves_avec_bulletin,
        }
        return render(request, 'Bull/dashboard.html', context)

    context = {}
    if user_role in ['admin', 'secretary']:
        nb_eleves = Student.objects.count()
        nb_enseignants = Teacher.objects.count()
        nb_classes = Classroom.objects.count()
        from Bull.models import User, SchoolYear, Term, Grade
        roles_stats = User.objects.values('role').annotate(count=Count('id'))
        school_year = SchoolYear.objects.filter(is_active=True).first()
        trimestres = Term.objects.filter(school_year=school_year) if school_year else []

        # Moyennes par classe/trimestre
        moyennes_classes = []
        if nb_classes > 0 and school_year:
            for classroom in Classroom.objects.all():
                for term in trimestres:
                    students = classroom.students.all()
                    total = 0
                    count = 0
                    for student in students:
                        avg = Grade.calculate_term_average(student, term)
                        if avg is not None:
                            total += avg
                            count += 1
                    moyenne = round(total / count, 2) if count > 0 else None
                    moyennes_classes.append({
                        'class': classroom.name,
                        'term': term.name,
                        'average': moyenne
                    })

        # Taux de réussite global
        total_students = Student.objects.count()
        reussite = 0
        if school_year:
            for student in Student.objects.all():
                avg = Grade.calculate_annual_average(student, school_year)
                if avg is not None and avg >= 10:
                    reussite += 1
        taux_reussite = round((reussite / total_students) * 100, 2) if total_students > 0 else None

        # Évolution des performances
        evolution = []
        for sy in SchoolYear.objects.all():
            total = 0
            count = 0
            for student in Student.objects.all():
                avg = Grade.calculate_annual_average(student, sy)
                if avg is not None:
                    total += avg
                    count += 1
            evolution.append({
                'year': sy.name,
                'average': round(total / count, 2) if count > 0 else None
            })

        # État de la BD
        from django.db import connection
        with connection.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table'")
            nb_tables = cursor.fetchone()[0]

        # Espace disque utilisé pour les bulletins
        
        bulletins_dir = os.path.join(os.path.dirname(__file__), '../bulletins')
        disk_usage = 0
        if os.path.exists(bulletins_dir):
            for root, dirs, files in os.walk(bulletins_dir):
                for f in files:
                    disk_usage += os.path.getsize(os.path.join(root, f))

        # Paramètres de l'application
        active_schoolyear = SchoolYear.objects.filter(is_active=True).first()

        context = {
            'nb_eleves': nb_eleves,
            'nb_enseignants': nb_enseignants,
            'nb_classes': nb_classes,
            'roles_stats': roles_stats,
            'school_year': school_year,
            'trimestres': trimestres,
            'moyennes_classes': moyennes_classes,
            'taux_reussite': taux_reussite,
            'evolution': evolution,
            'nb_tables': nb_tables,
            'disk_usage': disk_usage,
            'active_schoolyear': active_schoolyear,
            'dashboard_title': 'Tableau de bord Administrateur/Secrétariat',
            'total_bulletins': total_bulletins,
            'matiere_stats': matiere_stats,
            'moyennes_par_classe': moyennes_par_classe,
            'eleves_avec_bulletin': eleves_avec_bulletin,
        }
        return render(request, 'Bull/dashboard.html', context)
    # Pour les autres rôles, tu peux ajouter la logique ici
    return render(request, 'Bull/dashboard.html', {'dashboard_title': 'Tableau de bord', 'message': 'Dashboard personnalisé à venir.'})


@login_required
def profile_view(request):
    return render(request, 'Bull/profile.html')


# ---------------------------------------------
# Fonctions sanctions
# ---------------------------------------------
@require_GET
def sanctions_table(request):
    from Bull.models import Sanction
    sanctions = Sanction.objects.all().order_by('min_heures_absence')
    html = render_to_string('Bull/_sanctions_table.html', {'sanctions': sanctions})
    return JsonResponse({'html': html})

@csrf_exempt
def add_sanction(request):
    if request.method == 'POST':
        from Bull.models import Sanction
        texte = request.POST.get('texte')
        min_heures_absence = request.POST.get('min_heures_absence')
        if not texte or not min_heures_absence:
            return JsonResponse({'success': False, 'message': "Champs requis manquants."}, status=400)
        if Sanction.objects.filter(min_heures_absence=min_heures_absence).exists():
            return JsonResponse({'success': False, 'message': "Une sanction avec ce nombre d'heures existe déjà !"}, status=400)
        sanction = Sanction.objects.create(texte=texte, min_heures_absence=min_heures_absence)
        return JsonResponse({'success': True, 'message': "Sanction bien ajoutée !", 'sanction': {'id': sanction.id, 'texte': sanction.texte, 'min_heures_absence': sanction.min_heures_absence}})
    return JsonResponse({'success': False, 'message': 'Requête invalide.'}, status=400)

@csrf_exempt
def edit_sanction(request):
    if request.method == 'POST':
        from Bull.models import Sanction
        sanction_id = request.POST.get('id')
        texte = request.POST.get('texte')
        min_heures_absence = request.POST.get('min_heures_absence')
        if not texte or not min_heures_absence or not sanction_id:
            return JsonResponse({'success': False, 'message': "Champs requis manquants."}, status=400)
        try:
            sanction = Sanction.objects.get(id=sanction_id)
        except Sanction.DoesNotExist:
            return JsonResponse({'success': False, 'message': "Sanction introuvable."}, status=404)
        if Sanction.objects.filter(min_heures_absence=min_heures_absence).exclude(id=sanction_id).exists():
            return JsonResponse({'success': False, 'message': "Une sanction avec ce nombre d'heures existe déjà !"}, status=400)
        sanction.texte = texte
        sanction.min_heures_absence = min_heures_absence
        sanction.save()
        return JsonResponse({'success': True, 'message': "Sanction bien modifiée !", 'sanction': {'id': sanction.id, 'texte': sanction.texte, 'min_heures_absence': sanction.min_heures_absence}})
    return JsonResponse({'success': False, 'message': 'Requête invalide.'}, status=400)

@csrf_exempt
def delete_sanction(request):
    if request.method == 'POST':
        from Bull.models import Sanction
        sanction_id = request.GET.get('id')
        if not sanction_id:
            return JsonResponse({'success': False, 'message': "ID manquant."}, status=400)
        try:
            sanction = Sanction.objects.get(id=sanction_id)
            sanction.delete()
            return JsonResponse({'success': True, 'message': "Sanction supprimée !"})
        except Sanction.DoesNotExist:
            return JsonResponse({'success': False, 'message': "Sanction introuvable."}, status=404)
    return JsonResponse({'success': False, 'message': 'Requête invalide.'}, status=400)

# ---------------------------------------------
# Fonctions diverses
# ---------------------------------------------

@login_required
@user_passes_test(is_admin_or_secretary)
def users_view(request):
    return render(request, 'Bull/users.html')

@login_required
@user_passes_test(lambda u: u.role in ['teacher', 'admin', 'secretary'])
def grades_view(request):
    return render(request, 'Bull/grades.html')

@login_required
@user_passes_test(lambda u: u.role in ['teacher', 'admin', 'secretary'])
def bulletins_view(request):
    return render(request, 'Bull/bulletins.html')

@login_required
@user_passes_test(lambda u: u.role in ['parent', 'student', 'admin', 'secretary'])
def my_bulletin_view(request):
    return render(request, 'Bull/my_bulletin.html')

# Onglet paramètres : gestion des années, trimestres, séquences

def is_admin_or_secretary(user):
    return user.is_authenticated and (getattr(user, 'role', None) in ['admin', 'secretary'])
from django.views.decorators.http import require_POST

@login_required
@user_passes_test(is_admin_or_secretary)
def parameters_view(request):
    schoolyears = SchoolYear.objects.prefetch_related('terms__sequences').all().order_by('-name')
    from Bull.models import Sanction
    sanctions = Sanction.objects.all().order_by('min_heures_absence')
    return render(request, 'Bull/parameters.html', {'schoolyears': schoolyears, 'sanctions': sanctions})

# Actions pour créer, modifier, supprimer, activer
@login_required
@user_passes_test(is_admin_or_secretary)
@require_POST
def set_active_schoolyear(request, sy_id):
    SchoolYear.objects.update(is_active=False)
    SchoolYear.objects.filter(id=sy_id).update(is_active=True)
    return redirect('parameters')

@login_required
@user_passes_test(is_admin_or_secretary)
@require_POST
def set_active_sequence(request, seq_id):
    Sequence.objects.update(active=False)
    Sequence.objects.filter(id=seq_id).update(active=True)
    return redirect('parameters')

# Les vues add/edit/delete pour SchoolYear, Term, Sequence sont à ajouter si non présentes
# Vue édition année scolaire
@login_required
@user_passes_test(is_admin_or_secretary)
def edit_schoolyear(request, sy_id):
    from Bull.models import SchoolYear
    schoolyear = SchoolYear.objects.get(id=sy_id)
    if request.method == 'POST':
        name = request.POST.get('name')
        if name:
            schoolyear.name = name
            schoolyear.save()
            return redirect('parameters')
    return render(request, 'Bull/edit_schoolyear.html', {'schoolyear': schoolyear})


# Vue suppression année scolaire
@login_required
@user_passes_test(is_admin_or_secretary)
def delete_schoolyear(request, sy_id):
    from Bull.models import SchoolYear
    schoolyear = SchoolYear.objects.get(id=sy_id)
    if request.method == 'POST':
        schoolyear.delete()
        return redirect('parameters')
    return render(request, 'Bull/delete_schoolyear.html', {'schoolyear': schoolyear})



# Vue création année scolaire
@login_required
@user_passes_test(is_admin_or_secretary)
def add_schoolyear(request):
    from Bull.models import SchoolYear
    if request.method == 'POST':
        name = request.POST.get('name')
        if name:
            SchoolYear.objects.create(name=name, is_active=False)
            return redirect('parameters')
    return render(request, 'Bull/add_schoolyear.html')


# Vue création trimestre
@login_required
@user_passes_test(is_admin_or_secretary)
def add_term(request):
    from Bull.models import SchoolYear, Term
    schoolyear_id = request.GET.get('schoolyear')
    schoolyear = SchoolYear.objects.get(id=schoolyear_id) if schoolyear_id else None
    if request.method == 'POST':
        name = request.POST.get('name')
        if name and schoolyear:
            Term.objects.create(name=name, school_year=schoolyear)
            return redirect('parameters')
    return render(request, 'Bull/add_term.html', {'schoolyear': schoolyear})

# Vue édition trimestre
@login_required
@user_passes_test(is_admin_or_secretary)
def edit_term(request, term_id):
    from Bull.models import Term
    term = Term.objects.get(id=term_id)
    if request.method == 'POST':
        name = request.POST.get('name')
        if name:
            term.name = name
            term.save()
            return redirect('parameters')
    return render(request, 'Bull/edit_term.html', {'term': term})

# Vue suppression trimestre
@login_required
@user_passes_test(is_admin_or_secretary)
def delete_term(request, term_id):
    from Bull.models import Term
    term = Term.objects.get(id=term_id)
    if request.method == 'POST':
        term.delete()
        return redirect('parameters')
    return render(request, 'Bull/delete_term.html', {'term': term})

# Vue édition séquence
@login_required
@user_passes_test(is_admin_or_secretary)
def edit_sequence(request, seq_id):
    from Bull.models import Sequence
    seq = Sequence.objects.get(id=seq_id)
    if request.method == 'POST':
        name = request.POST.get('name')
        if name:
            seq.name = name
            seq.save()
            return redirect('parameters')
    return render(request, 'Bull/edit_sequence.html', {'sequence': seq})

# Vue suppression séquence
@login_required
@user_passes_test(is_admin_or_secretary)
def delete_sequence(request, seq_id):
    from Bull.models import Sequence
    seq = Sequence.objects.get(id=seq_id)
    if request.method == 'POST':
        seq.delete()
        return redirect('parameters')
    return render(request, 'Bull/delete_sequence.html', {'sequence': seq})


# -------------------------------------------
# vue pour les templates 
# --------------------------------------------
@staff_member_required
def bulletin_template_list(request):
    templates = BulletinTemplate.objects.all().order_by('-created_at')
    return render(request, 'Bull/bulletin_template_list.html', {'templates': templates})

@staff_member_required
def bulletin_template_create(request):
    if request.method == 'POST':
        form = BulletinTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('bulletin_template_list')
    else:
        form = BulletinTemplateForm()
    return render(request, 'Bull/bulletin_template_form.html', {'form': form})

@staff_member_required
def bulletin_template_edit(request, template_id):
    template = get_object_or_404(BulletinTemplate, id=template_id)
    if request.method == 'POST':
        form = BulletinTemplateForm(request.POST, request.FILES, instance=template)
        if form.is_valid():
            form.save()
            return redirect('bulletin_template_list')
    else:
        form = BulletinTemplateForm(instance=template)
    return render(request, 'Bull/bulletin_template_form.html', {'form': form, 'template': template})